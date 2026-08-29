from pathlib import Path
import hashlib
import json
import shutil

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(r"E:\postdoc-work\ist-project")
RESULTS = PROJECT / "results"
OUTPUT = PROJECT / "manuscript" / "supplementary_files_publication"
TABLE_DIR = OUTPUT / "supplementary_tables_csv"
OUTPUT.mkdir(parents=True, exist_ok=True)
TABLE_DIR.mkdir(parents=True, exist_ok=True)

MANIFEST = RESULTS / "step86g_supplementary_table_manifest.csv"
QC = RESULTS / "step86g_supplementary_table_qc.csv"
DOCX_OUTPUT = OUTPUT / "Supplementary_Tables_S1_to_S27.docx"
INDEX_CSV = OUTPUT / "Supplementary_Table_Index.csv"
README = OUTPUT / "README_Supplementary_Files.txt"


TABLES = [
    (1, "Physicochemical differences between Active and Inactive peptides",
     "step24_physicochemical_statistics_manuscript.csv", None),
    (2, "Highly correlated pairs among traditional physicochemical descriptors",
     "step26_high_correlation_pairs.csv", None),
    (3, "Locked-test performance of the four traditional classifiers",
     "step35_traditional_model_comparison_manuscript.csv", None),
    (4, "Cross-validated XGBoost permutation importance for traditional descriptors",
     "step38_xgboost_permutation_importance.csv",
     ["rank", "feature", "mean_cv_AUROC_drop", "sd_cv_AUROC_drop", "min_fold_AUROC_drop", "max_fold_AUROC_drop"]),
    (5, "Traditional-model consensus errors on the locked test set",
     "step40_consistently_misclassified_peptides.csv",
     ["ID", "sequence", "class", "models_correct", "models_wrong", "agreement_category", "prediction_pattern_LR_SVM_RF_XGB"]),
    (6, "Fold-contained ESM-2 scaling and PCA verification",
     "step47_esm2_preprocessing_fold_details.csv",
     ["fold", "train_rows", "validation_rows", "scaler_fit_samples", "pca_components", "pca_fit_samples", "cumulative_training_explained_variance", "train_validation_index_overlap", "locked_test_index_overlap"]),
    (7, "Locked-test performance of the four ESM-2 classifiers",
     "step52_esm2_model_comparison_manuscript.csv", None),
    (8, "Paired bootstrap comparison of Traditional and ESM-2 representations",
     "step54_paired_bootstrap_summary.csv",
     ["classifier", "metric", "observed_delta", "ci_95_lower", "ci_95_upper", "ci_excludes_zero", "ci_relation_to_zero", "bootstrap_replicates"]),
    (9, "Probability shifts across Traditional-to-ESM-2 prediction transitions",
     "step57_probability_shift_summary.csv",
     ["classifier", "transition_label", "n", "mean_correct_class_probability_gain", "median_correct_class_probability_gain", "positive_gain_count", "negative_gain_count", "zero_gain_count"]),
    (10, "Consensus hard cases across all eight frozen models",
     "step59_consensus_hard_cases_manuscript.csv",
     ["rank", "ID", "sequence", "true_class", "traditional_wrong_count", "esm2_wrong_count", "total_wrong_count", "all_models_mean_true_class_probability", "mean_absolute_descriptor_z", "most_extreme_descriptor", "difficulty_category"]),
    (11, "Amino-acid composition and sequence-pattern summaries for hard cases",
     "step60_group_amino_acid_summary.csv",
     ["metric", "hard_n", "hard_mean", "hard_median", "reference_n", "reference_mean", "reference_median", "hard_minus_reference_mean", "hard_to_reference_mean_ratio"]),
    (12, "Nearest-neighbor evidence for the 15 consensus hard cases",
     "step61_hard_case_nearest_neighbors.csv",
     ["rank", "ID", "true_class", "total_wrong_count", "sequence_opposite_class_proximity_margin", "esm2_opposite_class_proximity_margin", "neighborhood_pattern"]),
    (13, "Hard-case sequence-cluster sensitivity across similarity thresholds",
     "step62_cluster_threshold_summary.csv", None),
    (14, "Development-neighborhood purity and balanced-margin summaries",
     "step64_neighborhood_summary.csv", None),
    (15, "Frozen-model performance after progressive sequence-novelty filtering",
     "step67_sequence_novelty_performance.csv",
     ["model", "representation", "subset_label", "n", "active", "inactive", "AUROC", "AUPRC", "MCC", "F1"]),
    (16, "Error rates and confidence by development-similarity stratum",
     "step68_similarity_stratum_model_performance.csv",
     ["model", "representation", "similarity_stratum_label", "n", "active", "inactive", "wrong_count", "error_rate", "mean_true_class_probability"]),
    (17, "Probability-calibration metrics for all eight frozen models",
     "step71_calibration_metrics.csv",
     ["model", "brier_score", "brier_ci_lower", "brier_ci_upper", "log_loss", "log_loss_ci_lower", "log_loss_ci_upper", "ece_10_equal_width_bins", "ece_ci_lower", "ece_ci_upper", "calibration_intercept", "calibration_slope"]),
    (18, "Paired calibration comparison of Traditional and ESM-2 models",
     "step72_paired_calibration_comparison.csv",
     ["classifier", "metric", "traditional_point", "esm2_point", "observed_delta_esm2_minus_traditional", "ci_2_5", "ci_97_5", "ci_excludes_zero", "interval_direction"]),
    (19, "Decision-curve net-benefit summaries for all eight frozen models",
     "step73_decision_curve_model_summary.csv",
     ["model", "mean_net_benefit_0_05_to_0_50", "mean_net_benefit_0_05_to_0_20", "thresholds_beating_both", "maximum_observed_net_benefit", "threshold_at_maximum_observed_net_benefit"]),
    (20, "Stratified-bootstrap performance estimates and uncertainty",
     "step74_model_performance_bootstrap_summary.csv",
     ["model", "AUROC", "AUROC_CI_low", "AUROC_CI_high", "AUPRC", "AUPRC_CI_low", "AUPRC_CI_high", "MCC", "MCC_CI_low", "MCC_CI_high", "F1", "F1_CI_low", "F1_CI_high"]),
    (21, "Cross-validated canonical correlations between feature spaces",
     "step78_cv_cca_dimension_summary.csv", None),
    (22, "Traditional, ESM-2, and feature-fusion model comparison",
     "step79_fusion_model_comparison.csv",
     ["model", "classifier", "representation", "AUROC", "AUPRC", "MCC", "F1", "Accuracy", "Precision", "Recall", "Specificity"]),
    (23, "Paired bootstrap comparison of Fusion and ESM-2-only models",
     "step80_fusion_vs_esm2_paired_bootstrap_summary.csv",
     ["classifier", "metric", "esm2_point", "fusion_point", "observed_delta_fusion_minus_esm2", "ci_2_5", "ci_97_5", "ci_excludes_zero", "interval_status"]),
    (24, "Cross-model overlap of stable ESM-2 latent dimensions",
     "step81_esm2_feature_importance_overlap.csv", None),
    (25, "Peptide-level residue-perturbation sensitivity summary",
     "step82_peptide_perturbation_summary.csv", None),
    (26, "Residue-category sensitivity in hard and consensus-correct peptides",
     "step83_residue_category_summary.csv", None),
    (27, "Recurrent-motif context and residue sensitivity",
     "step83_motif_sensitivity_summary.csv", None),
]


def sha256(path):
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest().upper()


def safe_stem(title):
    text = "".join(ch if ch.isalnum() else "_" for ch in title)
    return "_".join(part for part in text.split("_") if part)[:80]


def humanize(name):
    special = {
        "ID": "ID", "AUROC": "AUROC", "AUPRC": "AUPRC", "MCC": "MCC",
        "F1": "F1", "ESM2": "ESM-2", "CV": "CV", "CI": "CI",
        "RF": "RF", "XGBoost": "XGBoost", "LR": "LR", "SVM": "SVM",
    }
    words = str(name).replace("esm2", "ESM2").replace("xgboost", "XGBoost").split("_")
    return " ".join(special.get(word, word.upper() if word in {"id", "cv", "ci"} else word.capitalize()) for word in words)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def display_value(value):
    if pd.isna(value):
        return ""
    if isinstance(value, (np.integer, int)):
        return str(int(value))
    if isinstance(value, (np.floating, float)):
        if abs(float(value)) >= 1000:
            return f"{float(value):,.1f}"
        if float(value).is_integer():
            return str(int(value))
        return f"{float(value):.3f}"
    text = str(value)
    if text in {"True", "False"}:
        return text
    return text


def column_widths(df, total_inches=9.8):
    weights = []
    for column in df.columns:
        samples = [humanize(column)] + [display_value(v) for v in df[column].head(40)]
        longest = max(len(value) for value in samples)
        weights.append(min(30, max(7, longest)))
    total = sum(weights)
    return [total_inches * weight / total for weight in weights]


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in (
        ("Title", 24, "0B2545", 0, 8),
        ("Heading 1", 16, "1F4D78", 14, 7),
        ("Heading 2", 12, "1F4D78", 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_data_table(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    widths = column_widths(df)
    header = table.rows[0]
    set_repeat_header(header)
    for index, column in enumerate(df.columns):
        cell = header.cells[index]
        cell.width = Inches(widths[index])
        cell.text = humanize(column)
        shade_cell(cell, "D9E7F2")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(7)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string("0B2545")
    for row_index, (_, record) in enumerate(df.iterrows(), start=1):
        cells = table.add_row().cells
        for col_index, column in enumerate(df.columns):
            cell = cells[col_index]
            cell.width = Inches(widths[col_index])
            cell.text = display_value(record[column])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 0:
                shade_cell(cell, "F4F7FA")
            for paragraph in cell.paragraphs:
                numeric = pd.api.types.is_numeric_dtype(df[column])
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if numeric else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7)
    return table


print("=" * 100)
print("STEP 86G - PUBLICATION SUPPLEMENTARY-TABLE PACKAGE")
print("=" * 100)

manifest_rows = []
qc_rows = []
loaded = []
for number, title, source_name, display_columns in TABLES:
    source = RESULTS / source_name
    if not source.exists():
        raise FileNotFoundError(source)
    df = pd.read_csv(source)
    if display_columns is not None:
        missing = [column for column in display_columns if column not in df.columns]
        if missing:
            raise ValueError(f"Table S{number}: missing display columns {missing}")
        display_df = df[display_columns].copy()
    else:
        display_df = df.copy()
    destination = TABLE_DIR / f"Table_S{number:02d}_{safe_stem(title)}.csv"
    shutil.copy2(source, destination)
    source_hash = sha256(source)
    destination_hash = sha256(destination)
    manifest_rows.append({
        "supplementary_table": f"Table S{number}", "title": title,
        "source_result_file": str(source), "publication_csv": str(destination),
        "source_rows": len(df), "source_columns": len(df.columns),
        "document_display_columns": len(display_df.columns),
        "source_sha256": source_hash, "publication_csv_sha256": destination_hash,
        "source_copy_exact": source_hash == destination_hash,
    })
    qc_rows.append({
        "supplementary_table": f"Table S{number}", "source_exists": source.exists(),
        "source_copy_exact": source_hash == destination_hash,
        "rows_match": len(pd.read_csv(destination)) == len(df),
        "columns_match": list(pd.read_csv(destination).columns) == list(df.columns),
        "document_values_from_source_only": True,
    })
    loaded.append((number, title, source_name, df, display_df))

pd.DataFrame(manifest_rows).to_csv(MANIFEST, index=False)
pd.DataFrame(qc_rows).to_csv(QC, index=False)
pd.DataFrame([{key: row[key] for key in ("supplementary_table", "title", "source_result_file", "publication_csv")}
              for row in manifest_rows]).to_csv(INDEX_CSV, index=False)

doc = Document()
configure_styles(doc)
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)
section.header.paragraphs[0].text = "Supplementary Tables"
section.header.paragraphs[0].style = doc.styles["Normal"]
section.header.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("687386")
add_page_number(section.footer.paragraphs[0])

title_p = doc.add_paragraph(style="Title")
title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_p.add_run("Supplementary Tables")
subtitle = doc.add_paragraph()
subtitle.add_run("Traditional physicochemical descriptors and frozen ESM-2 representations for anticancer-peptide classification").bold = True
meta = doc.add_paragraph()
meta.add_run("Submission-ready package | 27 sequential tables | Generated from validated project result files").italic = True
note = doc.add_paragraph()
note.add_run("Data policy. ").bold = True
note.add_run("Each numbered CSV is an exact byte-for-byte copy of its validated source result. The Word document presents reader-facing columns with rounded display values; the CSV files retain full precision and every original column.")

doc.add_heading("Contents", level=1)
index_table = doc.add_table(rows=1, cols=3)
index_table.style = "Table Grid"
index_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell, text in zip(index_table.rows[0].cells, ["Table", "Title", "Source result"]):
    cell.text = text
    shade_cell(cell, "D9E7F2")
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8)
for number, title, source_name, _, _ in loaded:
    cells = index_table.add_row().cells
    for cell, value in zip(cells, [f"Table S{number}", title, source_name]):
        cell.text = value
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.font.name = "Arial"
            run.font.size = Pt(7.5)

for number, title, source_name, full_df, display_df in loaded:
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Inches(11)
    new_section.page_height = Inches(8.5)
    new_section.top_margin = Inches(0.55)
    new_section.bottom_margin = Inches(0.55)
    new_section.left_margin = Inches(0.55)
    new_section.right_margin = Inches(0.55)
    new_section.header_distance = Inches(0.25)
    new_section.footer_distance = Inches(0.25)
    new_section.header.is_linked_to_previous = True
    new_section.footer.is_linked_to_previous = True

    heading = doc.add_paragraph(style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    heading.add_run(f"Table S{number}. {title}")
    source_note = doc.add_paragraph()
    source_note.paragraph_format.space_before = Pt(0)
    source_note.paragraph_format.space_after = Pt(5)
    source_note.paragraph_format.keep_with_next = True
    run = source_note.add_run(f"Source: results\\{source_name}. ")
    run.italic = True
    source_note.add_run(f"Full publication CSV: {full_df.shape[0]} rows x {full_df.shape[1]} columns. Word display: {display_df.shape[1]} selected reader-facing columns.")
    add_data_table(doc, display_df)

core = doc.core_properties
core.title = "Supplementary Tables S1-S27"
core.subject = "Publication supplementary tables generated from validated project results"
core.author = "ACP-ESM2 project"
core.keywords = "supplementary tables, anticancer peptides, ESM-2, machine learning"
doc.save(DOCX_OUTPUT)

readme_text = """PUBLICATION SUPPLEMENTARY FILES

Contents
--------
Supplementary_Tables_S1_to_S27.docx
    Manuscript-formatted Word document containing Tables S1-S27.

Supplementary_Table_Index.csv
    Table number, title, original result source, and publication CSV path.

supplementary_tables_csv\\Table_S01_... through Table_S27_...
    Exact copies of validated source result files, sequentially renamed for submission.

Data integrity
--------------
The numbered CSV files preserve full precision and all source columns. No model,
prediction, statistic, threshold, confidence interval, or interpretation was
recalculated. Reader-facing values in the Word document are rounded only for display.

Large bootstrap replicate files, raw prediction tables, environment logs, and internal
QC tables remain in the project results folder as reproducibility artifacts and are not
duplicated in this reader-facing submission package.
"""
README.write_text(readme_text, encoding="utf-8")

assert len(manifest_rows) == 27
assert all(row["source_copy_exact"] for row in qc_rows)
assert all(row["rows_match"] and row["columns_match"] for row in qc_rows)
assert DOCX_OUTPUT.exists() and DOCX_OUTPUT.stat().st_size > 0
assert INDEX_CSV.exists() and len(pd.read_csv(INDEX_CSV)) == 27

print(f"Supplementary tables packaged: {len(manifest_rows)}")
print(f"Exact source CSV copies: {sum(row['source_copy_exact'] for row in qc_rows)}/27")
print(f"Word document: {DOCX_OUTPUT}")
print(f"Output folder: {OUTPUT}")
print("STEP 86G COMPLETED SUCCESSFULLY")
